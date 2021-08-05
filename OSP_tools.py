#--------------------------------------------------------
#   LibUnits.py - Define system of units OpenSees model
#   Code:      II
#--------------------------------------------------------
# base units (SI units)
m = 1
kg = 1
s = 1
# other units
N = kg*m/s**2
cm = 0.01*m
Pa = N/m**2
ksi = 6894757.2932*Pa
kgf = 9.8066*N
MPa = 10**6*Pa
psi = 6894.76*Pa
# physical constants
g = 9.80665*m/s**2

# Propiedades de los materiales
fc = 210 # kg/cm2
E = 15100*fc**0.5*10**4*9.80665*Pa
G = 0.5*E/(1+0.2)
# Sección de Columna
a = 60*cm
Ac = a**2
ρlc = 2400*Ac*m**2
Izc = a**4/12
Iyc = a**4/12
Jxxc = 2.25*(a/2)**4
# Sección de Viga
b = 60*cm
h = 30*cm
A = b*h
ρl = 2400*A*m**2
Iz = b*h**3/12
Iy = b**3*h/12
Jxx = 0.229*max(b,h)*min(b,h)**3 # modificar
#
def GeoModel(dx, dy, h, nx, ny, nz):
    from numpy import zeros, ones
    import matplotlib.pyplot as plt
    # import matplotlib.pyplot as plt
    Lx, Ly, Lz = dx*nx, dy*ny, h*nz
    NN = (nx+1)*(ny+1)*(nz+1)
    Nodes = zeros((NN,5))
    # Creando los nodos y asignando coordenadas
    c = 0
    for i in range(nz+1):
        for j in range(ny+1):
            for k in range(nx+1):
                if k == nx and j != ny and j!= 0:
                    Nodes[c] = [c,k*dx,j*dy,i*h,0.50]
                elif k != nx and j == ny  and k!= 0:
                    Nodes[c] = [c,k*dx,j*dy,i*h,0.50]
                elif k == 0 and j != ny and j!= 0:
                    Nodes[c] = [c,k*dx,j*dy,i*h,0.50]
                elif k != nx and j == 0  and k!= 0:
                    Nodes[c] = [c,k*dx,j*dy,i*h,0.50]
                elif k == nx and j == ny:
                    Nodes[c] = [c,k*dx,j*dy,i*h,0.25]
                elif k == 0 and j == 0:
                    Nodes[c] = [c,k*dx,j*dy,i*h,0.25]
                elif k == nx and j == 0:
                    Nodes[c] = [c,k*dx,j*dy,i*h,0.25]
                elif k == 0 and j == ny:
                    Nodes[c] = [c,k*dx,j*dy,i*h,0.25]
                else:
                    Nodes[c] = [c,k*dx,j*dy,i*h,1.00]
                c = c + 1
    Nodes[:(nx+1)*(ny+1),4]=0
    # print(Nodes)

    NE = (nx*(ny+1)+ny*(nx+1)+(nx+1)*(ny+1))*nz
    Elems = zeros((NE,4))
    # Creando las conexiones de los elementos verticales
    c = 0
    for i in range(nz):
        for j in range(ny+1):
            for k in range(nx+1):
                Elems[c] = [c,c,c+(nx+1)*(ny+1),1]
                c = c + 1
    # Creando las conexiones de los elementos horizontales
    m = (nx+1)*(ny+1)
    for i in range(nz):
        for j in range(ny+1):
            for k in range(nx):
                Elems[c] = [c,m,m+1,2]
                m = m + 1
                c = c + 1
            m = m + 1
    # Creando las conexiones de los elementos horizontales
    n = 0 
    for i in range(nz):
        n = n + (nx+1)*(ny+1)
        for j in range(nx+1):
            for k in range(ny):
                Elems[c] = [c,j+k*(nx+1)+n,j+nx+1+k*(nx+1)+n,2]
                c = c + 1
    # Creando centro de diafragmas
    Diap = zeros((nz,4))
    for i in range(nz):
        Diap[i] = [i+1000,Lx/2.0,Ly/2.0,h*(i+1)]
    #
    return Nodes, Elems, Diap

def espectro_E030(T,Z=0.45,U=1.5,S=1.0,Tp=0.4,Tl=2.5,R=1):
    from numpy import zeros
    n = len(T)
    E030 = zeros(n)
    for i in range(n):
        if T[i]>=0 and T[i]<0.2*Tp:
            E030[i]=2.5#1+7.5*T[i]/Tp
        elif T[i]>=0.2*Tp and T[i]<Tp:
            E030[i]=2.5
        elif T[i]>=Tp and T[i]<Tl:
            E030[i] = 2.5*(Tp/T[i])
        elif T[i]>=Tl:
            E030[i] = 2.5*(Tp*Tl/T[i]**2)
        else:
            print("El periodo no puede ser negativo!")
    return E030*Z*U*S/R

def get_static_loads(coef,p,h,T):
    from numpy import zeros
    n = len(h)
    V = coef*sum(p)
    F = zeros(n)
    #
    if T > 0.0 and T <= 0.5:
        k=1.0
    elif T>0.5:
        k = 0.75+0.5*T
    else:
        print('El periodo es negativo!')
    #
    div = 0.
    for i in range(n):
        div = div + p[i]*h[i]**k
    #
    for i in range(n):
        F[i] = p[i]*h[i]**k/div*V
    return F,k

def getCombo(E030,MF,modo,NT,Tmodes):
    import numpy as np
    import pandas as pd
    # Obtenemos las Masas totales
    Mx = sum(sum(MF[0::3,0::3]))
    My = sum(sum(MF[1::3,1::3]))
    Mr = sum(sum(MF[2::3,2::3]))

    # Definimos valores iniciales
    Ux,Uy,Rz = np.zeros(NT),np.zeros(NT),np.zeros(NT)
    Ux[0::3]=1
    Uy[1::3]=1
    Rz[2::3]=1
    SUMx, SUMy, SUMr = 0., 0., 0.
    Nmodes = len(modo) 

    # Obtención de Masas Participativas
    ni=0
    np.set_printoptions(precision = 4)
    df1 = pd.DataFrame(columns=['Modo','T(s)','SumUx','SumUy','SumRz'])
    for j in range(1,Nmodes+1):
        FPx=modo[j-1].T@MF@Ux
        FPy=modo[j-1].T@MF@Uy
        FPr=modo[j-1].T@MF@Rz
        FPRx=FPx**2/Mx
        FPRy=FPy**2/My
        FPRr=FPr**2/Mr
        SUMx = SUMx + FPRx
        SUMy = SUMy + FPRy
        SUMr = SUMr + FPRr
        #
        if min(SUMx,SUMy,SUMr)>0.90 and ni==0:
            ni = j
        df1 = df1.append({'Modo':j, 'T(s)':Tmodes[j-1],'SumUx':SUMx,
                        'SumUy':SUMy,'SumRz':SUMr}, ignore_index=True)
    print('N° mínimo de Modos a considerar:',ni)

    # Definimos valores iniciales
    D_ABSx,D_RCSCx = np.zeros(NT),np.zeros(NT)
    Δ_ABSx,Δ_RCSCx = np.zeros(NT),np.zeros(NT)
    V_ABSx,V_RCSCx = np.zeros(NT),np.zeros(NT)
    D_ABSy,D_RCSCy = np.zeros(NT),np.zeros(NT)
    Δ_ABSy,Δ_RCSCy = np.zeros(NT),np.zeros(NT)
    V_ABSy,V_RCSCy = np.zeros(NT),np.zeros(NT)

    # Se realiza la Superpocisión Modal Espectral
    for j in range(1,ni+1):#ni+1
        FPx=modo[j-1].T@MF@Ux
        FPy=modo[j-1].T@MF@Uy
        FPr=modo[j-1].T@MF@Rz
        #
        Sa = E030[j-1]
        Sd = Sa*9.80665/(2*np.pi/Tmodes[j-1])**2
        #
        respDX = Sd*FPx*modo[j-1]
        respAX = Sa*FPx*MF@modo[j-1]
        D_ABSx = D_ABSx + abs(respDX)
        D_RCSCx = D_RCSCx + (respDX)**2
        respDX[3:] = respDX[3:] - respDX[:-3]
        Δ_ABSx = Δ_ABSx + abs(respDX)
        Δ_RCSCx = Δ_RCSCx + (respDX)**2
        V_ABSx = V_ABSx + abs(np.cumsum(respAX[::-1])[::-1])
        V_RCSCx = V_RCSCx + (np.cumsum(respAX[::-1])[::-1])**2
        #
        respDY = Sd*FPy*modo[j-1]
        respAY = Sa*FPy*MF@modo[j-1]
        D_ABSy = D_ABSy + abs(respDY)
        D_RCSCy = D_RCSCy + (respDY)**2
        respDY[3:] = respDY[3:] - respDY[:-3]
        Δ_ABSy = Δ_ABSy + abs(respDY)
        Δ_RCSCy = Δ_RCSCy + (respDY)**2
        V_ABSy = V_ABSy + abs(np.cumsum(respAY[::-1])[::-1])
        V_RCSCy = V_RCSCy + (np.cumsum(respAY[::-1])[::-1])**2

    # Se realiza la combinación 25%ABS + 75%RCSC
    D_RCSCx = D_RCSCx**0.5
    Δ_RCSCx = Δ_RCSCx**0.5
    V_RCSCx = V_RCSCx**0.5
    DDx = 0.25*D_ABSx + 0.75*D_RCSCx
    ΔDx = 0.25*Δ_ABSx + 0.75*Δ_RCSCx
    VDx = 0.25*V_ABSx + 0.75*V_RCSCx
    #
    D_RCSCy = D_RCSCy**0.5
    Δ_RCSCy = Δ_RCSCy**0.5
    V_RCSCy = V_RCSCy**0.5
    DDy = 0.25*D_ABSy + 0.75*D_RCSCy
    ΔDy = 0.25*Δ_ABSy + 0.75*Δ_RCSCy
    VDy = 0.25*V_ABSy + 0.75*V_RCSCy
    
    df2 = pd.DataFrame(columns=['Nivel','VDx(ton)','VDy(ton)','UDx(cm)','UDy(cm)'])
    for i in range(int(NT/3)):
        df2 = df2.append({'Nivel':i+1, 'VDx(ton)':VDx[0::3][i]/1000,
        'VDy(ton)':VDy[1::3][i]/1000,'UDx(cm)':DDx[0::3][i]*100,
        'UDy(cm)':DDy[1::3][i]*100}, ignore_index=True)

    return DDx, ΔDx, VDx, DDy, ΔDy, VDy, df1.iloc[:ni,:], df2

def genReport(df1,df2,df3,df4,df5,texto1,texto2):
    from PIL import Image
    import glob
    #
    lista = glob.glob('./imagenes/Mod*.png')
    #
    for archivo in lista:
        im = Image.open(archivo)
        width, height = im.size
        left, top = width/6, height/6
        right, bottom = 9*width/10, 9*height/10
        im1 = im.crop((left, top, right, bottom))
        im1.save(archivo)
    #
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    title1 = document.add_heading('Informe del Análisis Sísmico', 0)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph('Realizado por ')
    p.add_run('JPI Ingeniería e Innovación SAC').bold = True
    p.add_run(' para el curso ')
    p.add_run('ASEP.').italic = True

    document.add_paragraph('Edificio Analizado - vista 3D:')
    document.add_picture('./imagenes/Modelo_3D.png', width=Inches(5.0))
    document.add_paragraph('Edificación de Categoría Tipo C.')

    document.add_heading('Generalidades', level=1)
    document.add_paragraph('Metrado de Cargas', style='Intense Quote')

    document.add_paragraph('Para el metrado de cargas se consideró las siguientes cargas distribuidas:')
    document.add_paragraph('Carga Viva:\t\t\t\t250 kg/m2', style='List Bullet')
    document.add_paragraph('Carga de Losa:\t\t\t300 kg/m2', style='List Bullet')
    document.add_paragraph('Carga de Acabados:\t\t100 kg/m2', style='List Bullet')
    document.add_paragraph('Carga de Tabiquería:\t\t150 kg/m2', style='List Bullet')

    document.add_picture('./imagenes/Modelo_numerico.png', width=Inches(5.0))
    f1 = document.add_paragraph('Figura 1: Modelo Numérico para el Análisis.')
    f1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('Análisis Modal', level=1)
    document.add_paragraph('Modos de Vibración', style='Intense Quote')
    #
    t1 = document.add_paragraph('\nTabla 1: Factor de Participación de Masas.')
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table1 = document.add_table(rows=df1.shape[0]+1, cols=df1.shape[1])
    table1.style = 'Light Grid Accent 1'
    for j in range(df1.shape[-1]):
        table1.cell(0,j).text = df1.columns[j]
    for i in range(df1.shape[0]):
        for j in range(df1.shape[-1]):
            table1.cell(i+1,j).text = str(df1.values[i,j].round(4))

    document.add_picture('./imagenes/Modo_1.png', width=Inches(5.0))
    f2 = document.add_paragraph('Figura 2: Primer modo de vibración.')
    f2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_picture('./imagenes/Modo_2.png', width=Inches(5.0))
    f3 = document.add_paragraph('Figura 3: Segundo modo de vibración.')
    f3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_picture('./imagenes/Modo_3.png', width=Inches(5.0))
    f4 = document.add_paragraph('Figura 4: Tercer modo de vibración.')
    f4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #
    document.add_heading('Análisis Sísmico', level=1)
    document.add_paragraph('Análisis Estático', style='Intense Quote')
    p1 = document.add_paragraph(texto1)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #
    t2 = document.add_paragraph('Tabla 2: Fuerzas y desplazamientos del análisis estático en X.')
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table3 = document.add_table(rows=df3.shape[0]+1, cols=df3.shape[1])
    table3.style = 'Light Grid Accent 1'
    for j in range(df3.shape[-1]):
        table3.cell(0,j).text = df3.columns[j]
    for i in range(df3.shape[0]):
        for j in range(df3.shape[-1]):
            table3.cell(i+1,j).text = str(df3.values[i,j].round(4))
    t3 = document.add_paragraph('\nTabla 3: Fuerzas y desplazamientos del análisis estático en Y.')
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #
    table4 = document.add_table(rows=df4.shape[0]+1, cols=df4.shape[1])
    table4.style = 'Light Grid Accent 1'
    for j in range(df4.shape[-1]):
        table4.cell(0,j).text = df4.columns[j]
    for i in range(df4.shape[0]):
        for j in range(df4.shape[-1]):
            table4.cell(i+1,j).text = str(df4.values[i,j].round(4))
    #
    document.add_paragraph('Análisis Dinámico Modal Espectral', style='Intense Quote')
    #
    document.add_paragraph('En este análisis se consideraron los siguientes parámetros sísmicos:')
    document.add_paragraph('Factor de Zona:\t\t\t\tZ = 0.45', style='List Bullet')
    document.add_paragraph('Factor de Uso:\t\t\t\tU = 1.00', style='List Bullet')
    document.add_paragraph('F. de Amplificación del Suelo:\t\tS = 1.00', style='List Bullet')
    document.add_paragraph('Coef. de Reducción:\t\t\tRo= 8.00', style='List Bullet')

    document.add_picture('./imagenes/Espectro_E030.png', width=Inches(5.4))
    f5 = document.add_paragraph('Figura 5: Espectro según la norma E030.')
    f5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    t4 = document.add_paragraph('\nTabla 4: Respuesta Dinámica sin escalar.')
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table2 = document.add_table(rows=df2.shape[0]+1, cols=df2.shape[1])
    table2.style = 'Light Grid Accent 1'
    for j in range(df2.shape[-1]):
        table2.cell(0,j).text = df2.columns[j]
    for i in range(df2.shape[0]):
        for j in range(df2.shape[-1]):
            table2.cell(i+1,j).text = str(df2.values[i,j].round(4))
    #
    p2 = document.add_paragraph(texto2)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #
    t5 = document.add_paragraph('\nTabla 5: Respuesta Dinámica Escalada.')
    t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table5 = document.add_table(rows=df5.shape[0]+1, cols=df5.shape[1])
    table5.style = 'Light Grid Accent 1'
    for j in range(df5.shape[-1]):
        table5.cell(0,j).text = df5.columns[j]
    for i in range(df5.shape[0]):
        for j in range(df5.shape[-1]):
            table5.cell(i+1,j).text = str(df5.values[i,j].round(4))
    #
    document.add_page_break()
    document.add_heading('Resultados', level=1)
    document.add_paragraph('Distorsiones de Entrepiso', style='Intense Quote')
    document.add_picture('./imagenes/distorsion_din.png', width=Inches(5.0))
    f6 = document.add_paragraph('Figura 6: Distorsión de entrepiso del análisis dinámico.')
    f6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    file = 'Informe de Analisis Sismico.docx'
    document.save(file)
    import os
    os.startfile('%s'%file)